import os
import requests
from bs4 import BeautifulSoup
from docx import Document

def crawl_sina_news():
    url = "https://search.sina.com.cn/"
    params = {
        "q": "菜鸟集团",
        "range": "all",
        "c": "news",
        "sort": "rel",
        "num": 10,
        "page": 1
    }
    headers = {"User-Agent": "Mozilla/5.0"}

    response = requests.get(url, params=params, headers=headers)
    response.encoding = 'utf-8'
    soup = BeautifulSoup(response.text, 'html.parser')

    news_items = []
    # 新浪搜索结果的新闻条目在class="box-result clearfix"
    for item in soup.select('.box-result.clearfix'):
        title_tag = item.find('a')
        if title_tag:
            title = title_tag.get_text(strip=True)
            link = title_tag.get('href')
            news_items.append((title, link))
        if len(news_items) >= 10:
            break

    return news_items

def save_to_word(news_list, filename):
    doc = Document()
    doc.add_heading('菜鸟集团相关新闻 - 新浪新闻搜索', level=1)

    for idx, (title, link) in enumerate(news_list, 1):
        doc.add_paragraph(f"{idx}. {title}")
        doc.add_paragraph(f"链接: {link}")
        doc.add_paragraph("")

    doc.save(filename)
    print(f"已保存到 {filename}")

if __name__ == "__main__":
    news = crawl_sina_news()
    if not news:
        print("没有抓取到新闻内容，请检查网络或页面结构。")
    else:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        filepath = os.path.join(desktop, "菜鸟.docx")
        save_to_word(news, filepath)